"""
Clinic Schedule Parser — extracts client data from a Clinic Schedule.docx file.

The schedule contains a table mapping Client Ref → appointment details,
plus header-level metadata (clinic date, expert name, agency contact info).

Returns a lookup dict keyed by Client Ref (e.g. W2587243).
"""
import re
from pathlib import Path
from typing import Optional
from rich.console import Console

console = Console()


def parse_clinic_schedule(file_path: str) -> dict:
    """
    Parse a Clinic Schedule.docx and return a lookup dict keyed by Client Ref.

    Args:
        file_path: Path to the Clinic Schedule.docx file

    Returns:
        Dict keyed by client ref, e.g.:
        {
            "W2587243": {
                "appointment_time": "09:15",
                "client_name": "Mr Carl Hayes",
                "medical_records_required": False,
                "medical_agency": "SPEED",
                "clinic_date": "15/04/2026",
                "expert_name": "Dr Muhammad Jatoi",
            },
            ...
        }
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX support.\n"
            "Run: pip install python-docx"
        )

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Clinic schedule not found: {file_path}")

    doc = Document(file_path)

    # ── Extract header-level metadata from paragraphs ──────────────
    header_data = _extract_header_data(doc)

    # ── Extract table rows ─────────────────────────────────────────
    schedule = _extract_table_data(doc, header_data)

    console.print(
        f"  [green]✓[/green] Parsed clinic schedule: "
        f"{len(schedule)} client entries found"
    )

    return schedule


def _extract_header_data(doc) -> dict:
    """
    Extract metadata from the document paragraphs AND header tables.

    Looks for:
    - Expert name (from the doctor address block)
    - Clinic date
    - Agency contact info (phone, fax, email)
    """
    header = {
        "expert_name": None,
        "clinic_date": None,
        "agency_phone": None,
        "agency_email": None,
    }

    # Combine paragraph text AND table cell text for comprehensive search
    para_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_text += "\n" + cell.text

    full_text = para_text + "\n" + table_text

    # ── Expert name: "Dr <Name>" pattern in address block ──
    # Match only the name on a single line, not the full address
    dr_match = re.search(
        r'^(Dr\.?\s+[A-Z][a-z]+(?:[ \t]+[A-Z][a-z]+)+)', full_text, re.MULTILINE
    )
    if dr_match:
        header["expert_name"] = dr_match.group(1).strip()

    # ── Clinic Date: "Clinic Date" followed by a date in same or next cell ──
    # Check table cells first (more reliable — it's in a 2-cell row)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            for j, cell_text in enumerate(cells):
                if re.match(r'clinic\s*date', cell_text, re.IGNORECASE):
                    # Date might be in the same cell or the next cell
                    date_text = None
                    # Check same cell for inline date
                    inline = re.search(
                        r'clinic\s*date\s+(\d{1,2}\s+\w+\s+\d{4})',
                        cell_text, re.IGNORECASE
                    )
                    if inline:
                        date_text = inline.group(1)
                    # Check next cell
                    elif j + 1 < len(cells) and cells[j + 1]:
                        date_text = cells[j + 1]
                    if date_text:
                        header["clinic_date"] = _normalize_date(date_text)

    # Fallback: search paragraphs for clinic date
    if not header["clinic_date"]:
        date_match = re.search(
            r'Clinic\s+Date\s+(\d{1,2}\s+\w+\s+\d{4})', full_text, re.IGNORECASE
        )
        if date_match:
            header["clinic_date"] = _normalize_date(date_match.group(1))

    # ── Agency email ──
    email_match = re.search(r'Email:\s*(\S+@\S+)', full_text, re.IGNORECASE)
    if email_match:
        header["agency_email"] = email_match.group(1).strip()

    # ── Agency phone ──
    phone_match = re.search(r'Tel:\s*([\d\s]+)', full_text, re.IGNORECASE)
    if phone_match:
        header["agency_phone"] = re.sub(r'\s+', ' ', phone_match.group(1)).strip()

    return header


def _extract_table_data(doc, header_data: dict) -> dict:
    """
    Extract rows from the clinic schedule table.

    Expected columns (order may vary):
        Appt. Time | Client Ref. | Client Name | Medical Records | Medical Agency

    Note: The DOCX may contain multiple tables with a Client Ref column
    (e.g. an Attendance sheet). We identify the MAIN schedule table by
    requiring BOTH 'client_ref' AND 'medical_agency' columns to be present.
    Once found, we stop — so later tables don't overwrite the results.
    """
    schedule = {}

    for table in doc.tables:
        # Detect header row to find column indices
        if not table.rows:
            continue

        header_row = table.rows[0]
        header_cells = [cell.text.strip().lower() for cell in header_row.cells]

        # Build column index map
        col_map = _build_column_map(header_cells)

        # Require BOTH client_ref AND medical_agency to identify the
        # main schedule table (the Attendance table lacks medical_agency)
        if not col_map.get("client_ref") or not col_map.get("medical_agency"):
            continue

        # Process data rows
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]

            if len(cells) < 2:
                continue

            client_ref = _safe_get(cells, col_map.get("client_ref"))
            if not client_ref or not client_ref.startswith("W"):
                continue

            # Parse client name
            client_name = _safe_get(cells, col_map.get("client_name"))

            # Parse medical records
            med_records_text = _safe_get(cells, col_map.get("medical_records"))
            records_required = _parse_records_required(med_records_text)

            # Parse appointment time
            appt_time = _safe_get(cells, col_map.get("appt_time"))
            appt_time = _normalize_time(appt_time) if appt_time else None

            # Parse medical agency
            medical_agency = _safe_get(cells, col_map.get("medical_agency"))

            schedule[client_ref] = {
                "appointment_time": appt_time,
                "client_name": client_name,
                "medical_records_required": records_required,
                "medical_agency": medical_agency,
                "clinic_date": header_data.get("clinic_date"),
                "expert_name": header_data.get("expert_name"),
            }

        # Found and processed the main schedule table — stop here
        break

    return schedule


def _build_column_map(header_cells: list) -> dict:
    """
    Map semantic column names to their indices.
    Handles variations in column header text.
    """
    col_map = {}

    for i, cell in enumerate(header_cells):
        cell_lower = cell.lower().strip()

        if "time" in cell_lower and ("appt" in cell_lower or "app" in cell_lower):
            col_map["appt_time"] = i
        elif "client ref" in cell_lower or "ref" in cell_lower:
            if "client_ref" not in col_map:  # Prioritize "client ref" over just "ref"
                col_map["client_ref"] = i
        elif "client name" in cell_lower or "name" in cell_lower:
            if "client_name" not in col_map:
                col_map["client_name"] = i
        elif "medical record" in cell_lower or "record" in cell_lower:
            col_map["medical_records"] = i
        elif "medical agency" in cell_lower or "agency" in cell_lower:
            col_map["medical_agency"] = i

    return col_map


def _safe_get(cells: list, index: Optional[int]) -> Optional[str]:
    """Safely get a cell value by index."""
    if index is None or index >= len(cells):
        return None
    val = cells[index].strip()
    return val if val else None


def _parse_records_required(text: Optional[str]) -> Optional[bool]:
    """Parse the Medical Records column value."""
    if not text:
        return None

    lowered = text.lower().strip()
    if "not required" in lowered or "no" in lowered:
        return False
    elif "required" in lowered or "yes" in lowered:
        return True
    return None


def _normalize_time(time_str: str) -> Optional[str]:
    """Normalize time to HH:MM 24-hour format."""
    if not time_str:
        return None

    time_str = time_str.strip()

    # Already in HH:MM format
    match = re.match(r'^(\d{1,2}):(\d{2})$', time_str)
    if match:
        h, m = int(match.group(1)), int(match.group(2))
        return f"{h:02d}:{m:02d}"

    # Handle AM/PM
    match = re.match(r'^(\d{1,2}):(\d{2})\s*(AM|PM)$', time_str, re.IGNORECASE)
    if match:
        h, m = int(match.group(1)), int(match.group(2))
        period = match.group(3).upper()
        if period == "PM" and h != 12:
            h += 12
        elif period == "AM" and h == 12:
            h = 0
        return f"{h:02d}:{m:02d}"

    return time_str


def _normalize_date(date_str: str) -> Optional[str]:
    """Normalize a date string to dd/mm/yyyy."""
    if not date_str:
        return None

    from datetime import datetime

    # Remove ordinal suffixes
    cleaned = re.sub(r'(\d+)(st|nd|rd|th)', r'\1', date_str)
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()

    formats = [
        '%d %B %Y',    # 15 April 2026
        '%d %b %Y',    # 15 Apr 2026
        '%d/%m/%Y',    # 15/04/2026
        '%B %d %Y',    # April 15 2026
        '%Y-%m-%d',    # 2026-04-15
    ]

    for fmt in formats:
        try:
            dt = datetime.strptime(cleaned, fmt)
            return dt.strftime('%d/%m/%Y')
        except ValueError:
            continue

    return date_str


def find_clinic_schedule(folder_path: str) -> Optional[str]:
    """
    Search for a Clinic Schedule file in the given folder.
    Returns the path if found, None otherwise.

    Matches: 'Clinic Schedule.docx', 'clinic_schedule.docx', etc.
    """
    folder = Path(folder_path)

    for file in folder.iterdir():
        if not file.is_file():
            continue
        if file.name.startswith("~"):  # Skip Word temp files
            continue

        # Case-insensitive match for "clinic schedule" in the filename
        if "clinic" in file.name.lower() and "schedule" in file.name.lower():
            if file.suffix.lower() in {".docx", ".doc"}:
                return str(file)

    return None
