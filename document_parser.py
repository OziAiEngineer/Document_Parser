"""Document parser for multiple file formats (TXT, PDF, DOCX)."""
from pathlib import Path
from typing import Optional
import re


class DocumentParser:
    """Parse documents from various formats into plain text."""
    
    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.docx', '.doc'}
    
    def __init__(self):
        self._pdf_available = self._check_pdf_support()
        self._docx_available = self._check_docx_support()
    
    def _check_pdf_support(self) -> bool:
        """Check if PDF parsing libraries are available."""
        try:
            import pypdf
            return True
        except ImportError:
            try:
                import PyPDF2
                return True
            except ImportError:
                return False
    
    def _check_docx_support(self) -> bool:
        """Check if DOCX parsing library is available."""
        try:
            import docx
            return True
        except ImportError:
            return False
    
    def parse(self, file_path: Path) -> str:
        """
        Parse document and return text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            ImportError: If required library is not installed
        """
        extension = file_path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {extension}. "
                f"Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        if extension == '.txt':
            return self._parse_txt(file_path)
        elif extension == '.pdf':
            return self._parse_pdf(file_path)
        elif extension in {'.docx', '.doc'}:
            return self._parse_docx(file_path)
    
    def _parse_txt(self, file_path: Path) -> str:
        """Parse plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _parse_pdf(self, file_path: Path) -> str:
        """Parse PDF file using available library."""
        if not self._pdf_available:
            raise ImportError(
                "PDF parsing requires 'pypdf' library. "
                "Install with: pip install pypdf"
            )
        
        try:
            # Try pypdf first (newer)
            import pypdf
            return self._parse_pdf_pypdf(file_path)
        except ImportError:
            # Fall back to PyPDF2
            import PyPDF2
            return self._parse_pdf_pypdf2(file_path)
    
    def _parse_pdf_pypdf(self, file_path: Path) -> str:
        """Parse PDF using pypdf library."""
        import pypdf
        
        text_content = []
        
        with open(file_path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            
            if getattr(reader, 'is_encrypted', False):
                if not self._decrypt_pdf(reader, file_path):
                    raise Exception("File has not been decrypted")
            
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
                except Exception as e:
                    print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
        
        full_text = '\n\n'.join(text_content)
        return self._clean_extracted_text(full_text)
    
    def _parse_pdf_pypdf2(self, file_path: Path) -> str:
        """Parse PDF using PyPDF2 library."""
        import PyPDF2
        
        text_content = []
        
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            
            if getattr(reader, 'is_encrypted', False):
                if not self._decrypt_pdf(reader, file_path):
                    raise Exception("File has not been decrypted")
            
            for page_num in range(len(reader.pages)):
                try:
                    page = reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
                except Exception as e:
                    print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
        
        full_text = '\n\n'.join(text_content)
        return self._clean_extracted_text(full_text)
    
    def _decrypt_pdf(self, reader, file_path: Path) -> bool:
        """Attempt to decrypt the PDF using known passwords and directory context."""
        if not getattr(reader, 'is_encrypted', False):
            return True
            
        # 1. Check for specific condition: Admiral Law Limited
        for doc_file in file_path.parent.glob("*.doc*"):
            if doc_file == file_path:
                continue
            try:
                doc_text = ""
                try:
                    doc_text = self._parse_docx(doc_file)
                except Exception:
                    pass
                
                if "Admiral Law Limited" in doc_text:
                    if reader.decrypt("Med1calR3port?"):
                        return True
            except Exception:
                pass
                
        # 2. Try other known passwords
        known_passwords = ["smes", "Med1calR3port?"]
        for pwd in known_passwords:
            try:
                if reader.decrypt(pwd):
                    return True
            except Exception:
                pass
                
        return False

    def _parse_docx(self, file_path: Path) -> str:
        """Parse DOCX file."""
        if not self._docx_available:
            raise ImportError(
                "DOCX parsing requires 'python-docx' library. "
                "Install with: pip install python-docx"
            )
        
        from docx import Document
        
        try:
            doc = Document(file_path)
            
            text_content = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_content.append(row_text)
            
            full_text = '\n\n'.join(text_content)
            return self._clean_extracted_text(full_text)
            
        except Exception as e:
            raise Exception(f"Failed to parse DOCX file: {e}")
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean extracted text from PDF/DOCX artifacts."""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove form feed and other control characters
        text = re.sub(r'[\x0c\x0b]', '\n', text)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive spaces
        text = re.sub(r' +', ' ', text)
        
        return text.strip()
    
    def get_supported_formats(self) -> dict:
        """Return dictionary of supported formats and their availability."""
        return {
            'txt': True,
            'pdf': self._pdf_available,
            'docx': self._docx_available
        }
