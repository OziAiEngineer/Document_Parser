"""Utility functions for text processing and date normalization."""
import re
import os
import subprocess
from pathlib import Path
from datetime import datetime


# ──────────────────────────────────────────────
# DATE NORMALIZATION
# ──────────────────────────────────────────────

def normalize_date(date_str: str) -> str:
    """
    Normalize various date formats to dd/mm/yyyy.

    Handles formats like:
    - "3rd March 2023"
    - "03-03-2023"
    - "March 3rd 2023"
    - "14/02/2023"
    - "14 February 2023"
    - "2023-03-14"
    - "14th of February 2023"
    """
    if not date_str or not isinstance(date_str, str):
        return None

    original = date_str.strip()

    # Remove "of" e.g. "14th of February 2023" → "14th February 2023"
    date_str = re.sub(r'\bof\b', '', original, flags=re.IGNORECASE)

    # Remove ordinal suffixes (1st, 2nd, 3rd, 4th...)
    date_str = re.sub(r'(\d+)(st|nd|rd|th)', r'\1', date_str)

    # Collapse extra whitespace after removals
    date_str = re.sub(r'\s+', ' ', date_str).strip()

    # Common date formats to try
    formats = [
        '%d/%m/%Y',
        '%d-%m-%Y',
        '%d.%m.%Y',
        '%d %B %Y',       # 14 February 2023
        '%d %b %Y',       # 14 Feb 2023
        '%B %d %Y',       # February 14 2023
        '%b %d %Y',       # Feb 14 2023
        '%Y-%m-%d',       # 2023-03-14
        '%Y/%m/%d',       # 2023/03/14
        '%d %m %Y',       # 14 02 2023
        '%m/%d/%Y',       # 03/14/2023 (US format fallback)
    ]

    for fmt in formats:
        try:
            dt = datetime.strptime(date_str.strip(), fmt)
            return dt.strftime('%d/%m/%Y')
        except ValueError:
            continue

    # Last resort — return original so we don't silently lose a date
    return original


# ──────────────────────────────────────────────
# TEXT CLEANING
# ──────────────────────────────────────────────

def clean_text(text: str) -> str:
    """Clean and normalize text content."""
    if not text:
        return ""

    # Normalize unicode whitespace characters
    text = text.replace('\xa0', ' ').replace('\t', ' ')

    # Remove extra whitespace but preserve paragraph breaks
    text = re.sub(r' {2,}', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


def clean_phone(phone: str) -> str:
    """
    Normalize a phone number string.
    Returns null if the number is unavailable, missing, or invalid.

    Examples:
    - "07928524454"       → "07928524454"
    - "Unavailable"       → None
    - "N/A"               → None
    - "+44 7928 524454"   → "+447928524454"
    """
    if not phone or not isinstance(phone, str):
        return None

    lowered = phone.strip().lower()

    # Treat these as missing
    null_values = {"unavailable", "n/a", "na", "none", "not available", "-", "unknown", ""}
    if lowered in null_values:
        return None

    # Remove spaces and dashes from valid numbers
    cleaned = re.sub(r'[\s\-\(\)]', '', phone.strip())

    # Must contain at least 7 digits to be valid
    if len(re.sub(r'\D', '', cleaned)) < 7:
        return None

    return cleaned


# ──────────────────────────────────────────────
# POSTCODE EXTRACTION
# ──────────────────────────────────────────────

def extract_postcode(text: str) -> str:
    """Extract UK postcode from text."""
    # Full UK postcode pattern
    pattern = r'\b([A-Z]{1,2}\d{1,2}[A-Z]?\s?\d[A-Z]{2})\b'
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(0).upper() if match else None


def extract_address_without_postcode(address: str) -> str:
    """
    Strip postcode from an address string.
    Useful when address and postcode are in the same line.

    Example:
    "11 Kelvin House, London, SE26 5FB" → "11 Kelvin House, London"
    """
    if not address:
        return None

    postcode_pattern = r'\b[A-Z]{1,2}\d{1,2}[A-Z]?\s?\d[A-Z]{2}\b'
    cleaned = re.sub(postcode_pattern, '', address, flags=re.IGNORECASE)
    return re.sub(r',?\s*$', '', cleaned.strip())


# ──────────────────────────────────────────────
# DOCUMENT READING
# ──────────────────────────────────────────────

def read_document(file_path: str) -> str:
    """
    Read any supported document format and return plain text.

    Supported formats:
    - .txt  → built-in
    - .pdf  → pypdf
    - .docx → python-docx
    - .doc  → pywin32 (Windows) or LibreOffice (Linux/Mac)
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")

    suffix = path.suffix.lower()

    handlers = {
        ".txt":  _read_txt,
        ".pdf":  _read_pdf,
        ".docx": _read_docx,
        ".doc":  _read_doc,
    }

    if suffix not in handlers:
        _raise_unsupported_format(suffix)

    return clean_text(handlers[suffix](file_path))


def _read_txt(file_path: str) -> str:
    """Read plain text file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def _decrypt_pdf(reader, file_path_str: str) -> bool:
    """Attempt to decrypt the PDF using known passwords and directory context."""
    if not getattr(reader, 'is_encrypted', False):
        return True
        
    file_path = Path(file_path_str)
    
    # 1. Check for specific condition: Admiral Law Limited
    for doc_file in file_path.parent.glob("*.doc*"):
        if str(doc_file) == file_path_str:
            continue
        try:
            doc_text = ""
            try:
                doc_text = read_document(str(doc_file))
            except Exception:
                pass
            
            if "Admiral Law Limited" in doc_text:
                if reader.decrypt("Med1calR3port?"):
                    return True
        except Exception:
            pass
            
    # 2. Try other known passwords
    known_passwords = ["smes", "Med1calR3port?"]
    for pwd in known_passwords:
        try:
            if reader.decrypt(pwd):
                return True
        except Exception:
            pass
            
    return False


def _read_pdf(file_path: str) -> str:
    """Read PDF using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf is required for PDF support.\nRun: pip install pypdf")

    reader = PdfReader(file_path)
    
    if getattr(reader, 'is_encrypted', False):
        if not _decrypt_pdf(reader, file_path):
            raise Exception("File has not been decrypted")

    # Only attempt extraction if successfully decrypted or unencrypted
    pages = []
    for page in reader.pages:
        try:
            text = page.extract_text()
            if text:
                pages.append(text)
        except Exception:
            pass
            
    return "\n".join(pages)


def _read_docx(file_path: str) -> str:
    """Read modern .docx format using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX support.\nRun: pip install python-docx")

    doc = Document(file_path)

    # Extract paragraphs AND tables
    content = []
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                content.append(row_text)

    return "\n".join(content)


def _read_doc(file_path: str) -> str:
    """
    Read legacy .DOC format using multiple fallback methods.

    Priority:
    1. pywin32 (Windows + Word installed) — most accurate
    2. LibreOffice headless convert to txt
    3. LibreOffice headless convert to docx then read
    """

    # Method 1 — pywin32 (Windows only, requires MS Word)
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(file_path))
        text = doc.Content.Text
        doc.Close(False)
        word.Quit()
        return text
    except Exception:
        pass

    # Method 2 — LibreOffice convert to .txt
    try:
        output_dir = Path(file_path).parent
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "txt:Text",
             "--outdir", str(output_dir), file_path],
            check=True, capture_output=True, timeout=30
        )
        txt_path = Path(file_path).with_suffix(".txt")
        if txt_path.exists():
            return _read_txt(str(txt_path))
    except Exception:
        pass

    # Method 3 — LibreOffice convert to .docx then read
    try:
        output_dir = Path(file_path).parent
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx",
             "--outdir", str(output_dir), file_path],
            check=True, capture_output=True, timeout=30
        )
        docx_path = Path(file_path).with_suffix(".docx")
        if docx_path.exists():
            return _read_docx(str(docx_path))
    except Exception:
        pass

    raise RuntimeError(
        f"Could not read .DOC file: {file_path}\n"
        f"Try one of these fixes:\n"
        f"  1. pip install pywin32        (Windows with MS Word installed)\n"
        f"  2. Install LibreOffice and add it to your system PATH\n"
        f"  3. Manually open and Save As .docx then retry"
    )


def _raise_unsupported_format(suffix: str) -> None:
    """Raise a clear error showing supported formats."""
    supported = {
        "TXT":  ("✓ Available", "Built-in"),
        "PDF":  ("✓ Available", "pypdf          → pip install pypdf"),
        "DOCX": ("✓ Available", "python-docx    → pip install python-docx"),
        "DOC":  ("✓ Available", "pywin32        → pip install pywin32"),
    }

    table = f"\n{'Format':<8}{'Status':<16}{'Required Library'}\n"
    table += "─" * 55 + "\n"
    for fmt, (status, lib) in supported.items():
        table += f"{fmt:<8}{status:<16}{lib}\n"

    raise ValueError(
        f"Format '{suffix}' is not supported.\n\n"
        f"Supported Formats:{table}"
    )